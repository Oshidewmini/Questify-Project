import io
from typing import List
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

class PaperExportEngine:

    @staticmethod
    def generate_docx(title: str, exam_board: str, subject: str, questions: list, include_mark_scheme: bool = False) -> io.BytesIO:
        doc = Document()

        # Document Header
        title_p = doc.add_paragraph()
        title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_p.add_run(f"{exam_board.upper()} EXAMINATION")
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(30, 58, 138)  # Primary accent color

        sub_p = doc.add_paragraph()
        sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_p.add_run(f"Subject: {subject} | Paper Title: {title}")
        sub_run.font.size = Pt(12)
        sub_run.font.italic = True

        doc.add_paragraph("-" * 50)

        # Questions
        for idx, q in enumerate(questions, 1):
            qp = doc.add_paragraph()
            q_num = qp.add_run(f"Question {idx} [{q.get('bloom_level', 'Remember')}] ({q.get('mark_value', 1)} Mark(s)) - AO: {q.get('ao_code', 'AO1')}\n")
            q_num.bold = True
            
            q_text = qp.add_run(q.get("text", ""))
            q_text.font.size = Pt(11)

            if q.get("options"):
                for opt in q["options"]:
                    op_p = doc.add_paragraph(f"    {opt}")
                    op_p.paragraph_format.left_indent = Inches(0.5)

            if include_mark_scheme and q.get("answer"):
                ans_p = doc.add_paragraph()
                ans_p.paragraph_format.left_indent = Inches(0.25)
                ans_run = ans_p.add_run(f"Mark Scheme: {q.get('answer')}")
                ans_run.font.color.rgb = RGBColor(16, 185, 129)  # Green accent for answer key
                ans_run.font.italic = True

            doc.add_paragraph()  # Spacing

        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    @staticmethod
    def generate_pdf(title: str, exam_board: str, subject: str, questions: list, include_mark_scheme: bool = False) -> io.BytesIO:
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, rightMargin=36, leftMargin=36, topMargin=36, bottomMargin=36)
        story = []

        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'TitleStyle',
            parent=styles['Heading1'],
            fontSize=18,
            textColor=colors.HexColor('#1E3A8A'),
            alignment=1,
            spaceAfter=12
        )
        subtitle_style = ParagraphStyle(
            'SubTitleStyle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#4B5563'),
            alignment=1,
            spaceAfter=24
        )
        q_style = ParagraphStyle(
            'QuestionStyle',
            parent=styles['Normal'],
            fontSize=11,
            leading=14,
            spaceAfter=8
        )

        story.append(Paragraph(f"{exam_board.upper()} EXAMINATION", title_style))
        story.append(Paragraph(f"Subject: {subject} | Paper Title: {title}", subtitle_style))
        story.append(Spacer(1, 12))

        for idx, q in enumerate(questions, 1):
            q_header = f"<b>Q{idx}. [{q.get('bloom_level', 'Remember')}] ({q.get('mark_value', 1)} Marks) - {q.get('ao_code', 'AO1')}</b>"
            story.append(Paragraph(q_header, q_style))
            story.append(Paragraph(q.get("text", ""), q_style))

            if q.get("options"):
                for opt in q["options"]:
                    story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;{opt}", q_style))

            if include_mark_scheme and q.get("answer"):
                mark_text = f"<font color='#10B981'><b>Mark Scheme:</b> {q.get('answer')}</font>"
                story.append(Paragraph(mark_text, q_style))

            story.append(Spacer(1, 10))

        doc.build(story)
        buffer.seek(0)
        return buffer
