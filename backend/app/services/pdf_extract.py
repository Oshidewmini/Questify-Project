import io
from typing import List, Tuple

import pymupdf as fitz

TEXT_LENGTH_THRESHOLD = 20
RENDER_DPI = 300
OCR_LANG = "eng"
MAX_FILE_BYTES = 10 * 1024 * 1024


def _ocr_available() -> Tuple[bool, str]:
    try:
        import pytesseract
        pytesseract.get_tesseract_version()
        return True, ""
    except Exception as exc:
        return False, str(exc)


def _ocr_page(page, dpi: int = RENDER_DPI, lang: str = OCR_LANG) -> str:
    from PIL import Image, ImageOps
    import pytesseract

    zoom = dpi / 72
    pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom))
    img = Image.open(io.BytesIO(pix.tobytes("png")))
    img = ImageOps.autocontrast(img.convert("L"))
    return pytesseract.image_to_string(img, lang=lang).strip()


def extract_pdf_bytes(data: bytes) -> Tuple[str, List[str]]:
    warnings: List[str] = []
    ocr_ok, ocr_err = _ocr_available()
    if not ocr_ok:
        warnings.append(
            "OCR is unavailable (Tesseract not installed). Scanned pages may extract little text."
        )

    doc = fitz.open(stream=data, filetype="pdf")
    pages: List[str] = []
    try:
        for page_num in range(len(doc)):
            page = doc[page_num]
            direct_text = page.get_text().strip()
            if len(direct_text) >= TEXT_LENGTH_THRESHOLD:
                pages.append(direct_text)
                continue
            if ocr_ok:
                try:
                    pages.append(_ocr_page(page))
                except Exception as exc:
                    warnings.append(f"OCR failed on page {page_num + 1}: {exc}")
                    pages.append(direct_text)
            else:
                pages.append(direct_text)
                if not direct_text:
                    warnings.append(
                        f"Page {page_num + 1} had little extractable text and OCR was skipped."
                    )
    finally:
        doc.close()

    return "\n".join(pages), warnings


def extract_docx_bytes(data: bytes) -> Tuple[str, List[str]]:
    from docx import Document

    document = Document(io.BytesIO(data))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    text = "\n".join(parts)
    warnings = [] if text.strip() else ["No text could be extracted from this Word file."]
    return text, warnings


def extract_txt_bytes(data: bytes) -> Tuple[str, List[str]]:
    for encoding in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            text = data.decode(encoding)
            return text, []
        except UnicodeDecodeError:
            continue
    return data.decode("utf-8", errors="replace"), ["File was not valid UTF-8; replacement characters were used."]


def extract_file(filename: str, data: bytes) -> Tuple[str, List[str]]:
    if len(data) > MAX_FILE_BYTES:
        raise ValueError(f"{filename} exceeds the 10 MB limit.")

    name = (filename or "").lower()
    if name.endswith(".pdf"):
        return extract_pdf_bytes(data)
    if name.endswith(".docx"):
        return extract_docx_bytes(data)
    if name.endswith(".txt"):
        return extract_txt_bytes(data)
    raise ValueError(f"Unsupported file type: {filename}. Use PDF, DOCX, or TXT.")
