import io
import re
from typing import Any, Dict, List

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from .allocation import QUESTION_TYPES

SECTION_TITLES = {
    "MCQ": "Multiple Choice Questions",
    "True/False": "True / False",
    "Short Answer": "Short Answer Questions",
    "Fill-in": "Fill in the Blank",
    "Structured": "Structured Questions",
    "Essay": "Essay / Extended Response",
}

BLOOM_RANK = {
    "Remember": 0,
    "Understand": 1,
    "Apply": 2,
    "Analyze": 3,
    "Evaluate": 4,
    "Create": 5,
}

SECTION_LETTERS = "ABCDEFGHIJKLMNOPQRSTUVWXYZ"


def _sort_by_bloom(questions: List[dict]) -> List[dict]:
    return sorted(
        questions,
        key=lambda q: BLOOM_RANK.get(str(q.get("bloom_level") or ""), 99),
    )


def _safe_filename(title: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", title or "question_paper").strip("_")
    return (slug or "question_paper") + ".docx"


def _add_centered(doc: Document, text: str, *, bold: bool = False, italic: bool = False, size: int = 12):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return para


def _group_by_type(questions: List[dict]) -> Dict[str, List[dict]]:
    grouped: Dict[str, List[dict]] = {t: [] for t in QUESTION_TYPES}
    for q in questions:
        qtype = q.get("question_type") or "Short Answer"
        grouped.setdefault(qtype, []).append(q)
    return grouped


def _render_question_body(doc: Document, q: dict, number: int):
    qtype = q.get("question_type")
    marks = q.get("mark_value") or 1
    p = doc.add_paragraph()
    p.add_run(f"{number}. {q.get('text') or ''}  ")
    p.add_run(f"[{marks} mark(s)]").italic = True

    if qtype == "MCQ":
        options = q.get("options") or {}
        if isinstance(options, dict):
            for letter in ["A", "B", "C", "D"]:
                if letter in options:
                    doc.add_paragraph(f"    {letter}) {options[letter]}")
        elif isinstance(options, list):
            for idx, opt in enumerate(options):
                doc.add_paragraph(f"    {chr(65 + idx)}) {opt}")

    elif qtype == "True/False":
        doc.add_paragraph("    True / False")

    elif qtype == "Structured":
        for part in q.get("parts") or []:
            part_p = doc.add_paragraph()
            label = part.get("label") or ""
            part_p.add_run(f"    ({label}) {part.get('question') or ''}  ")
            part_p.add_run(f"[{part.get('marks', 1)} mark(s)]").italic = True

    doc.add_paragraph()


def _render_answer(doc: Document, q: dict, number: int):
    qtype = q.get("question_type")
    if qtype == "MCQ":
        doc.add_paragraph(f"{number}. {q.get('correct_option') or q.get('answer') or '?'}")
        return
    if qtype == "True/False":
        doc.add_paragraph(f"{number}. {q.get('correct_option') or q.get('answer') or '?'}")
        return
    if qtype == "Fill-in":
        doc.add_paragraph(f"{number}. {q.get('answer') or '?'}")
        return
    if qtype == "Structured":
        heading = doc.add_paragraph()
        heading.add_run(f"{number}. {q.get('text') or ''} [{q.get('mark_value', 0)} marks]").bold = True
        for part in q.get("parts") or []:
            doc.add_paragraph(f"    ({part.get('label')}) {part.get('question') or ''}")
            for point in part.get("marking_scheme") or []:
                doc.add_paragraph(f"        - {point}")
        doc.add_paragraph()
        return

    heading = doc.add_paragraph()
    heading.add_run(f"{number}. {q.get('text') or ''} [{q.get('mark_value', 0)} marks]").bold = True
    scheme = q.get("marking_scheme") or []
    if scheme:
        for point in scheme:
            doc.add_paragraph(f"    - {point}")
    elif q.get("answer"):
        doc.add_paragraph(f"    {q.get('answer')}")
    doc.add_paragraph()


def build_docx_bytes(header: Dict[str, Any], questions: List[dict]) -> bytes:
    doc = Document()
    grouped = _group_by_type(questions)
    present_types = [t for t in QUESTION_TYPES if grouped.get(t)]
    total_marks = sum(int(q.get("mark_value") or 0) for q in questions)

    _add_centered(doc, header.get("title") or "Question Paper", bold=True, size=18)
    _add_centered(
        doc,
        f"{header.get('exam_board', '')}  |  {header.get('qualification_level', '')}  |  {header.get('subject', '')}",
        italic=True,
        size=12,
    )
    duration = header.get("duration_minutes")
    duration_bit = f"   |   Duration: {duration} minutes" if duration else ""
    _add_centered(
        doc,
        f"Total Questions: {len(questions)}   |   Total Marks: {total_marks}{duration_bit}",
        size=11,
    )
    name_p = doc.add_paragraph()
    name_p.add_run("Name: ____________________________        Date: ______________")
    doc.add_paragraph("_" * 90)

    for idx, qtype in enumerate(present_types):
        letter = SECTION_LETTERS[idx]
        title = SECTION_TITLES.get(qtype, qtype)
        doc.add_heading(f"Section {letter}: {title}", level=1)
        ordered = _sort_by_bloom(grouped[qtype])
        for i, q in enumerate(ordered, 1):
            _render_question_body(doc, q, i)

    doc.add_page_break()
    doc.add_heading("Answer Key", level=1)
    for idx, qtype in enumerate(present_types):
        letter = SECTION_LETTERS[idx]
        title = SECTION_TITLES.get(qtype, qtype)
        doc.add_heading(f"Section {letter} — {title}", level=2)
        ordered = _sort_by_bloom(grouped[qtype])
        for i, q in enumerate(ordered, 1):
            _render_answer(doc, q, i)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


def filename_for(title: str) -> str:
    return _safe_filename(title)
